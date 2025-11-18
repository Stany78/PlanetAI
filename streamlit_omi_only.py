import io
from datetime import datetime

import pandas as pd
import streamlit as st

from agent_core import geocode_indirizzo
from omi_utils import get_quotazione_omi_da_coordinate, warmup_omi_cache

# ---------------------------------------------------------
# Word report (opzionale, se python-docx è installato)
# ---------------------------------------------------------
try:
    from docx import Document  # type: ignore

    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False


# ---------------------------------------------------------
# Utility
# ---------------------------------------------------------
def build_word_report(
    *,
    comune_input: str,
    indirizzo_input: str,
    lat: float,
    lon: float,
    zona_omi,
) -> io.BytesIO:
    """
    Crea un report Word basato sui soli dati OMI (€/m²),
    senza superficie e senza valore totale.
    Il report è strutturato in stile "mini perizia".
    """
    document = Document()

    # -------------------------------------------------
    # Titolo e intestazione
    # -------------------------------------------------
    document.add_heading("Report di valutazione OMI", level=1)
    document.add_paragraph(
        f"Data generazione report: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    document.add_paragraph(
        "Il presente documento riporta una stima sintetica basata esclusivamente "
        "sulle quotazioni OMI (Osservatorio del Mercato Immobiliare - Agenzia delle Entrate), "
        "espresse in €/m², per la zona in cui ricade l'indirizzo indicato."
    )

    # -------------------------------------------------
    # 1. Dati di input
    # -------------------------------------------------
    document.add_heading("1. Dati di input", level=2)

    p = document.add_paragraph()
    p.add_run("Comune inserito: ").bold = True
    p.add_run(comune_input)

    p = document.add_paragraph()
    p.add_run("Indirizzo inserito: ").bold = True
    p.add_run(indirizzo_input)

    p = document.add_paragraph()
    p.add_run("Coordinate geografiche (lat, lon): ").bold = True
    p.add_run(f"{lat:.6f}, {lon:.6f}")

    # -------------------------------------------------
    # 2. Inquadramento della zona OMI
    # -------------------------------------------------
    document.add_heading("2. Zona OMI di riferimento", level=2)

    comune_omi = str(getattr(zona_omi, "comune", ""))
    provincia_omi = str(getattr(zona_omi, "provincia", ""))
    codice_zona = str(getattr(zona_omi, "zona_codice", ""))
    descr_zona = str(getattr(zona_omi, "zona_descrizione", ""))

    p = document.add_paragraph()
    p.add_run("Comune (OMI): ").bold = True
    p.add_run(comune_omi)

    p = document.add_paragraph()
    p.add_run("Provincia: ").bold = True
    p.add_run(provincia_omi)

    p = document.add_paragraph()
    p.add_run("Zona OMI: ").bold = True
    p.add_run(codice_zona)

    p = document.add_paragraph()
    p.add_run("Descrizione zona: ").bold = True
    p.add_run(descr_zona)

    # Eventuali campi aggiuntivi se presenti nel dataclass (anno / semestre)
    anno = getattr(zona_omi, "anno", None)
    semestre = getattr(zona_omi, "semestre", None)
    if anno is not None and semestre is not None:
        p = document.add_paragraph()
        p.add_run("Periodo OMI di riferimento: ").bold = True
        p.add_run(f"{anno} – semestre {semestre}")

    # -------------------------------------------------
    # 3. Quotazioni OMI €/m² (compravendita)
    # -------------------------------------------------
    document.add_heading("3. Quotazioni OMI €/m² (compravendita)", level=2)

    val_min = float(zona_omi.val_min_mq)
    val_med = float(zona_omi.val_med_mq)
    val_max = float(zona_omi.val_max_mq)

    # Tabella valori
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    # Riga titolo sezione
    hdr = table.rows[0].cells
    hdr[0].text = "Parametro"
    hdr[1].text = "Valore"

    # Min
    row_min = table.rows[1].cells
    row_min[0].text = "Valore minimo €/m²"
    row_min[1].text = f"{val_min:,.0f} €/m²".replace(",", ".")

    # Med
    row_med = table.rows[2].cells
    row_med[0].text = "Valore mediano €/m²"
    row_med[1].text = f"{val_med:,.0f} €/m²".replace(",", ".")

    # Max
    row_max = table.rows[3].cells
    row_max[0].text = "Valore massimo €/m²"
    row_max[1].text = f"{val_max:,.0f} €/m²".replace(",", ".")

    # -------------------------------------------------
    # 4. Interpretazione sintetica
    # -------------------------------------------------
    document.add_heading("4. Interpretazione sintetica", level=2)

    # Classificazione qualitativa molto semplice basata sul valore mediano
    if val_med < 2000:
        fascia = "fascia tendenzialmente medio-bassa per il contesto urbano."
    elif 2000 <= val_med < 3500:
        fascia = "fascia mediamente in linea con i valori urbani di riferimento."
    elif 3500 <= val_med < 5000:
        fascia = "fascia medio-alta rispetto alla media urbana."
    else:
        fascia = "fascia alta, relativa ad ambiti particolarmente richiesti o centrali."

    document.add_paragraph(
        f"Sulla base del valore mediano pari a circa {val_med:,.0f} €/m² "
        f"(arrotondato), la zona OMI '{codice_zona}' può essere considerata in "
        f"{fascia}"
    )

    document.add_paragraph(
        "Il valore minimo rappresenta generalmente immobili con caratteristiche "
        "meno favorevoli (stato di manutenzione scadente, piano basso, esposizione "
        "penalizzata, contesto meno richiesto), mentre il valore massimo si riferisce "
        "a immobili con caratteristiche migliori (buona esposizione, piano alto, stato "
        "manutentivo buono/ottimo, contesti più pregiati)."
    )

    # -------------------------------------------------
    # 5. Limiti e note metodologiche
    # -------------------------------------------------
    document.add_heading("5. Limiti e note metodologiche", level=2)

    document.add_paragraph(
        "Le quotazioni OMI sono valori indicativi di zona, espressi in €/m², "
        "elaborati dall'Osservatorio del Mercato Immobiliare dell'Agenzia delle Entrate. "
        "Esse non tengono conto delle specifiche caratteristiche del singolo immobile, "
        "come stato manutentivo, piano, presenza di ascensore, spazi esterni, vista, "
        "anno di costruzione o ristrutturazione, qualità del condominio, ecc."
    )

    document.add_paragraph(
        "Il presente report non costituisce una perizia asseverata, ma uno strumento "
        "di supporto alla valutazione basato su dati statistici ufficiali di mercato."
    )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------
# Cache OMI
# ---------------------------------------------------------
@st.cache_resource(show_spinner="📦 Caricamento dati OMI (solo al primo avvio)...")
def init_omi():
    warmup_omi_cache()
    return True


init_omi()

# ---------------------------------------------------------
# UI Streamlit
# ---------------------------------------------------------
st.set_page_config(
    page_title="PlanetAI – Valutazione OMI",
    page_icon="🏙️",
    layout="centered",
)

st.title("🏙️ PlanetAI – Valutazione OMI")

st.write(
    """
Inserisci **Comune** e **Indirizzo** (via, civico).

L'app:
1. geocoda l'indirizzo,
2. trova la **zona OMI**,
3. mostra i valori OMI **min / med / max €/m²**,
4. permette di scaricare un **report Word** solo con quotazioni €/m².
"""
)

with st.form("omi_form"):
    col_a, col_b = st.columns(2)
    with col_a:
        comune = st.text_input("Comune", value="Como", placeholder="Es: Como")
    with col_b:
        indirizzo = st.text_input(
            "Indirizzo",
            value="Via Borgovico 150",
            placeholder="Es: Via Borgovico 150",
        )

    submit = st.form_submit_button("Calcola valutazione OMI 🧮")

if not submit:
    st.stop()

if not comune.strip() or not indirizzo.strip():
    st.error("Inserisci sia il **Comune** che l'**Indirizzo**.")
    st.stop()

with st.spinner("📍 Geocoding indirizzo e ricerca zona OMI..."):
    # 1) Geocoding
    lat, lon = geocode_indirizzo(comune, indirizzo)

    # 2) Quotazione OMI da coordinate
    zona_omi = get_quotazione_omi_da_coordinate(lat, lon)

if zona_omi is None or zona_omi.val_med_mq is None:
    st.error(
        "Non è stato possibile trovare una zona OMI per queste coordinate. "
        "Verifica che l'indirizzo sia corretto oppure che i dati OMI/KML coprano questa zona."
    )
    st.stop()

st.success("✅ Zona OMI trovata!")

# ---------------------------------------------------------
# Dettagli zona
# ---------------------------------------------------------
st.subheader("📌 Zona OMI")

zona_col1, zona_col2 = st.columns(2)
with zona_col1:
    st.write(f"**Comune (OMI):** {zona_omi.comune}")
    st.write(f"**Provincia:** {zona_omi.provincia}")
with zona_col2:
    st.write(f"**Zona OMI:** {zona_omi.zona_codice}")
    st.write(f"**Descrizione zona:** {zona_omi.zona_descrizione}")

st.markdown("---")

# ---------------------------------------------------------
# Valori €/m²
# ---------------------------------------------------------
st.subheader("💶 Valori OMI €/m²")

col_min, col_med, col_max = st.columns(3)
col_min.metric("Minimo", f"{zona_omi.val_min_mq:,.0f} €/m²".replace(",", "."))
col_med.metric("Mediano", f"{zona_omi.val_med_mq:,.0f} €/m²".replace(",", "."))
col_max.metric("Massimo", f"{zona_omi.val_max_mq:,.0f} €/m²".replace(",", "."))

# 🔁 Istogramma invertito: Massimo → Mediano → Minimo
df_valori = pd.DataFrame(
    {
        "Tipologia": ["Massimo", "Mediano", "Minimo"],
        "Valore €/m²": [
            zona_omi.val_max_mq,
            zona_omi.val_med_mq,
            zona_omi.val_min_mq,
        ],
    }
)
st.bar_chart(
    data=df_valori.set_index("Tipologia"),
    height=260,
)

st.caption("Fonte: dati OMI caricati dai file CSV e KML (Agenzia delle Entrate).")

st.markdown("---")

# ---------------------------------------------------------
# Download report Word
# ---------------------------------------------------------
if HAVE_DOCX:
    report_buffer = build_word_report(
        comune_input=comune,
        indirizzo_input=indirizzo,
        lat=lat,
        lon=lon,
        zona_omi=zona_omi,
    )

    file_name = (
        f"Report_OMI_{comune.replace(' ', '_')}_"
        f"{indirizzo.replace(' ', '_')}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    )

    st.download_button(
        label="📄 Scarica report Word (quotazioni €/m²)",
        data=report_buffer,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
else:
    st.info(
        "Per abilitare il download del report Word, aggiungi `python-docx` al file `requirements.txt`."
    )
