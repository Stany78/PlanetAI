"""
Generatore report Word combinato:
- Dati OMI (Agenzia Entrate)
- Dati Immobiliare.it
- Confronto e analisi
"""

import os
from datetime import datetime
from typing import Dict, Optional
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from map_generator import crea_mappa_interattiva


def genera_report_combinato(
    comune: str,
    via: str,
    lat: float,
    lon: float,
    raggio_km: float,
    zona_omi: Optional[Dict],
    stats_immobiliare: Optional[Dict],
    appartamenti: Optional[list] = None,
    analisi_ai: Optional[Dict] = None,
    output_dir: str = "reports"
) -> str:
    """
    Genera report Word combinato con dati OMI + Immobiliare.it
    
    Args:
        comune: Nome comune
        via: Nome via
        lat: Latitudine
        lon: Longitudine
        raggio_km: Raggio ricerca
        zona_omi: Dict con dati OMI (da omi_utils.OMIQuotazione)
        stats_immobiliare: Dict con statistiche Immobiliare.it
        appartamenti: Lista appartamenti (per mappa)
        output_dir: Directory output
    
    Returns:
        Path del file generato
    """
    
    print("\n📝 Generazione report Word combinato...")
    
    # Crea documento
    doc = Document()
    
    # Stile base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    now = datetime.now()
    
    # ========================================
    # TITOLO
    # ========================================
    title = doc.add_heading('REPORT ANALISI IMMOBILIARE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading('Dati OMI + Mercato Immobiliare.it', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # ========================================
    # SEZIONE 1: INFORMAZIONI RICERCA
    # ========================================
    doc.add_heading('📋 Informazioni Ricerca', 1)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ['Data estrazione', now.strftime('%d/%m/%Y %H:%M')],
        ['Località', f'{via}, {comune}'],
        ['Coordinate', f'lat: {lat:.6f}, lon: {lon:.6f}'],
        ['Raggio di ricerca', f'{raggio_km} km'],
        ['Fonte dati', 'OMI (Agenzia Entrate) + Immobiliare.it']
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # ========================================
    # SEZIONE 2: DATI OMI
    # ========================================
    doc.add_heading('🏛️ Dati OMI (Agenzia delle Entrate)', 1)
    
    if zona_omi and zona_omi.get('val_med_mq') is not None:
        doc.add_paragraph(f"Zona OMI identificata: {zona_omi['zona_codice']}")
        doc.add_paragraph(f"Descrizione: {zona_omi['zona_descrizione']}")
        doc.add_paragraph(f"Comune: {zona_omi['comune']} ({zona_omi['provincia']})")
        doc.add_paragraph()
        
        # Valori OMI
        doc.add_paragraph('Valori €/mq (dati ufficiali rogiti):', style='Heading 3')
        
        omi_table = doc.add_table(rows=2, cols=3)
        omi_table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['Minimo', 'Mediano', 'Massimo']
        for i, header in enumerate(headers):
            cell = omi_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Valori
        omi_table.rows[1].cells[0].text = f"€{zona_omi['val_min_mq']:,.0f}".replace(',', '.')
        omi_table.rows[1].cells[1].text = f"€{zona_omi['val_med_mq']:,.0f}".replace(',', '.')
        omi_table.rows[1].cells[2].text = f"€{zona_omi['val_max_mq']:,.0f}".replace(',', '.')
        
        for i in range(3):
            omi_table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('⚠️ Dati OMI non disponibili per questa zona.')
    
    doc.add_paragraph()
    
    # ========================================
    # SEZIONE 3: DATI IMMOBILIARE.IT
    # ========================================
    doc.add_heading('🏢 Analisi Mercato (Immobiliare.it)', 1)
    
    if stats_immobiliare and stats_immobiliare['n_appartamenti'] > 0:
        doc.add_paragraph(f"Appartamenti analizzati: {stats_immobiliare['n_appartamenti']}")
        doc.add_paragraph()
        
        # Statistiche Prezzi
        doc.add_paragraph('Statistiche Prezzi:', style='Heading 3')
        doc.add_paragraph(f"Prezzo medio: €{stats_immobiliare['prezzo']['medio']:,.0f}".replace(',', '.'))
        doc.add_paragraph(f"Prezzo mediano: €{stats_immobiliare['prezzo']['mediano']:,.0f}".replace(',', '.'))
        doc.add_paragraph(f"Prezzo minimo: €{stats_immobiliare['prezzo']['min']:,.0f}".replace(',', '.'))
        doc.add_paragraph(f"Prezzo massimo: €{stats_immobiliare['prezzo']['max']:,.0f}".replace(',', '.'))
        doc.add_paragraph()
        
        # Statistiche Superfici
        doc.add_paragraph('Statistiche Superfici:', style='Heading 3')
        doc.add_paragraph(f"Superficie media: {stats_immobiliare['mq']['medio']:.0f} m²")
        doc.add_paragraph(f"Superficie mediana: {stats_immobiliare['mq']['mediano']:.0f} m²")
        doc.add_paragraph(f"Superficie minima: {stats_immobiliare['mq']['min']:.0f} m²")
        doc.add_paragraph(f"Superficie massima: {stats_immobiliare['mq']['max']:.0f} m²")
        doc.add_paragraph()
        
        # Prezzo al MQ
        doc.add_paragraph('Prezzo al Metro Quadro:', style='Heading 3')
        doc.add_paragraph(f"Prezzo/mq medio: €{stats_immobiliare['prezzo_mq']['medio']:,.0f}/m²".replace(',', '.'))
        doc.add_paragraph(f"Prezzo/mq mediano: €{stats_immobiliare['prezzo_mq']['mediano']:,.0f}/m²".replace(',', '.'))
        doc.add_paragraph(f"Prezzo/mq minimo: €{stats_immobiliare['prezzo_mq']['min']:,.0f}/m²".replace(',', '.'))
        doc.add_paragraph(f"Prezzo/mq massimo: €{stats_immobiliare['prezzo_mq']['max']:,.0f}/m²".replace(',', '.'))
        doc.add_paragraph()
        
        # Distribuzione per fasce prezzo
        doc.add_paragraph('Distribuzione per fasce di prezzo:', style='Heading 3')
        df = stats_immobiliare['dataframe']
        
        fasce_prezzo = [
            ('Fino a €200.000', df[df['prezzo'] <= 200000].shape[0]),
            ('€200.000 - €350.000', df[(df['prezzo'] > 200000) & (df['prezzo'] <= 350000)].shape[0]),
            ('€350.000 - €500.000', df[(df['prezzo'] > 350000) & (df['prezzo'] <= 500000)].shape[0]),
            ('Oltre €500.000', df[df['prezzo'] > 500000].shape[0])
        ]
        
        for fascia, count in fasce_prezzo:
            doc.add_paragraph(f"  • {fascia}: {count} appartamenti", style='List Bullet')
        
        doc.add_paragraph()
        
        # Distribuzione per superfici
        doc.add_paragraph('Distribuzione per fasce di superficie:', style='Heading 3')
        
        fasce_mq = [
            ('Fino a 60 m²', df[df['mq'] <= 60].shape[0]),
            ('60 - 100 m²', df[(df['mq'] > 60) & (df['mq'] <= 100)].shape[0]),
            ('100 - 150 m²', df[(df['mq'] > 100) & (df['mq'] <= 150)].shape[0]),
            ('Oltre 150 m²', df[df['mq'] > 150].shape[0])
        ]
        
        for fascia, count in fasce_mq:
            doc.add_paragraph(f"  • {fascia}: {count} appartamenti", style='List Bullet')
        
        doc.add_paragraph()
        
    else:
        doc.add_paragraph('⚠️ Nessun appartamento trovato su Immobiliare.it per questa zona.')
    
    # ========================================
    # SEZIONE 4: ANALISI PER AGENZIA
    # ========================================
    if stats_immobiliare and stats_immobiliare['n_appartamenti'] > 0:
        doc.add_heading('🏢 Analisi per Agenzia', 1)
        
        # Prepara dati agenzie
        df = stats_immobiliare['dataframe']
        agenzie = df.groupby('agenzia').agg({
            'prezzo': ['count', 'mean'],
            'mq': 'mean'
        }).reset_index()
        
        agenzie.columns = ['Agenzia', 'N° Appartamenti', 'Prezzo Medio', 'MQ Medio']
        agenzie = agenzie.sort_values('N° Appartamenti', ascending=False)
        
        # Tabella agenzie
        table = doc.add_table(rows=len(agenzie)+1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header
        headers = ['Agenzia', 'N° Appartamenti', 'Prezzo Medio', 'MQ Medio']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dati
        for idx, row in agenzie.iterrows():
            table.rows[idx+1].cells[0].text = str(row['Agenzia'])
            table.rows[idx+1].cells[1].text = str(int(row['N° Appartamenti']))
            table.rows[idx+1].cells[2].text = f"€{row['Prezzo Medio']:,.0f}".replace(',', '.')
            table.rows[idx+1].cells[3].text = f"{row['MQ Medio']:.0f} m²"
            
            table.rows[idx+1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.rows[idx+1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.rows[idx+1].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
    
    # ========================================
    # SEZIONE 4.5: ANALISI DEVELOPER
    # ========================================
    if zona_omi and stats_immobiliare and stats_immobiliare.get('n_appartamenti', 0) > 0:
        doc.add_heading('💼 Analisi per Developer/Investitori', 1)
        
        n_app = stats_immobiliare['n_appartamenti']
        omi_med = zona_omi.get('val_med_mq')
        mercato_med = stats_immobiliare['prezzo_mq']['mediano']
        
        # 1. SATURAZIONE MERCATO
        doc.add_heading('1. Saturazione Mercato', 2)
        
        sat_table = doc.add_table(rows=1, cols=2)
        sat_table.style = 'Light Grid Accent 1'
        
        sat_data = [
            ['Appartamenti in Vendita (Nuove Costruzioni)', str(n_app)]
        ]
        
        for i, (label, value) in enumerate(sat_data):
            sat_table.rows[i].cells[0].text = label
            sat_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            sat_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Valutazione
        doc.add_paragraph('Valutazione:', style='Heading 3')
        if n_app < 10:
            doc.add_paragraph('✓ Mercato LIBERO - Poca concorrenza. Buona opportunità di ingresso.')
        elif n_app < 30:
            doc.add_paragraph('• Mercato MEDIO - Concorrenza normale. Necessaria differenziazione.')
        else:
            doc.add_paragraph('⚠ Mercato SATURO - Alta concorrenza. Necessaria forte differenziazione.')
        
        doc.add_paragraph()
        
        # 2. PRICING BENCHMARK
        doc.add_heading('2. Pricing Benchmark', 2)
        
        mercato_min = stats_immobiliare['prezzo_mq']['min']
        mercato_max = stats_immobiliare['prezzo_mq']['max']
        gap_percentuale = ((mercato_med - omi_med) / omi_med) * 100
        
        price_table = doc.add_table(rows=5, cols=2)
        price_table.style = 'Light Grid Accent 1'
        
        price_data = [
            ['OMI Mediano (Baseline)', f"€{omi_med:,.0f}/m²"],
            ['Mercato Range', f"€{mercato_min:,.0f} - €{mercato_max:,.0f}/m²"],
            ['Mercato Mediano', f"€{mercato_med:,.0f}/m²"],
            ['Gap vs OMI', f"{gap_percentuale:+.1f}%"],
            ['Sweet Spot Consigliato', f"€{mercato_med:,.0f}/m²"]
        ]
        
        for i, (label, value) in enumerate(price_data):
            price_table.rows[i].cells[0].text = label
            price_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            price_table.rows[i].cells[1].text = value.replace(',', '.')
        
        doc.add_paragraph()
        
        # 3. GAP ANALYSIS STRATEGICO
        doc.add_heading('3. Gap Analysis Strategico', 2)
        
        doc.add_paragraph('Interpretazione:', style='Heading 3')
        
        if gap_percentuale > 50:
            doc.add_paragraph('⚠ GAP MOLTO ALTO (+50% vs OMI)')
            doc.add_paragraph('• Mercato con forte premium su nuove costruzioni')
            doc.add_paragraph('• Possibile sopravvalutazione - Alto rischio pricing')
            doc.add_paragraph('• Raccomandazione: Verificare qualità, considera pricing conservativo')
        elif gap_percentuale > 30:
            doc.add_paragraph('⚠ GAP SIGNIFICATIVO (+30-50% vs OMI)')
            doc.add_paragraph('• Premium pricing per nuove costruzioni')
            doc.add_paragraph('• Mercato accetta sovrapprezzo elevato')
            doc.add_paragraph('• Raccomandazione: Giustifica premium con alta qualità e marketing forte')
        elif gap_percentuale > 15:
            doc.add_paragraph('✓ GAP NORMALE (+15-30% vs OMI)')
            doc.add_paragraph('• Premium standard nuove costruzioni')
            doc.add_paragraph('• Mercato equilibrato con margini sani')
            doc.add_paragraph('• Raccomandazione: Sweet spot ideale per sviluppo')
        else:
            doc.add_paragraph('• GAP BASSO (<15% vs OMI)')
            doc.add_paragraph('• Prezzi allineati a valori rogiti')
            doc.add_paragraph('• Mercato competitivo con margini contenuti')
            doc.add_paragraph('• Raccomandazione: Ottimizza costi, focus su volume vendite')
        
        doc.add_paragraph()
        
        # 4. ANALISI AGENZIE
        doc.add_heading('4. Principali Operatori', 2)
        
        if stats_immobiliare.get('dataframe') is not None:
            df = stats_immobiliare['dataframe']
            agenzie_stats = df.groupby('agenzia').size().reset_index(name='count')
            agenzie_stats = agenzie_stats.sort_values('count', ascending=False).head(5)
            
            # Tabella Top 5 agenzie
            ag_table = doc.add_table(rows=len(agenzie_stats) + 1, cols=3)
            ag_table.style = 'Light Grid Accent 1'
            
            # Header
            ag_table.rows[0].cells[0].text = 'Agenzia'
            ag_table.rows[0].cells[1].text = 'N° Appartamenti'
            ag_table.rows[0].cells[2].text = '% Mercato'
            for cell in ag_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Dati
            for i, (_, row) in enumerate(agenzie_stats.iterrows(), 1):
                percentuale = (row['count'] / n_app * 100)
                ag_table.rows[i].cells[0].text = row['agenzia']
                ag_table.rows[i].cells[1].text = str(int(row['count']))
                ag_table.rows[i].cells[2].text = f"{percentuale:.1f}%"
            
            doc.add_paragraph()
            
            # Concentrazione mercato
            top3_count = agenzie_stats.head(3)['count'].sum()
            top3_share = (top3_count / n_app * 100)
            
            doc.add_paragraph(f'Concentrazione Top 3: {top3_share:.1f}%', style='Heading 3')
            
            if top3_share > 60:
                doc.add_paragraph('⚠ Mercato concentrato - Pochi operatori dominanti')
            elif top3_share > 40:
                doc.add_paragraph('• Mercato moderato - Mix operatori grandi/piccoli')
            else:
                doc.add_paragraph('✓ Mercato frammentato - Molti piccoli operatori')
        
        doc.add_paragraph()
    
    # SEZIONE 5: CONFRONTO OMI vs MERCATO
    # ========================================
    doc.add_heading('📊 Confronto OMI vs Mercato', 1)
    
    if zona_omi and zona_omi.get('val_med_mq') and stats_immobiliare:
        omi_med = zona_omi['val_med_mq']
        mercato_med = stats_immobiliare['prezzo_mq']['mediano']
        
        gap = ((mercato_med - omi_med) / omi_med) * 100
        
        doc.add_paragraph(f"Valore OMI mediano: €{omi_med:,.0f}/m²".replace(',', '.'))
        doc.add_paragraph(f"Prezzo mercato mediano: €{mercato_med:,.0f}/m²".replace(',', '.'))
        doc.add_paragraph(f"Gap: {gap:+.1f}%".replace('.', ','))
        doc.add_paragraph()
        
        # Interpretazione
        doc.add_paragraph('Interpretazione:', style='Heading 3')
        if gap > 15:
            doc.add_paragraph('Il mercato quota significativamente sopra i valori OMI. Possibile zona ad alta domanda o prezzi di offerta ottimistici.')
        elif gap > 5:
            doc.add_paragraph('Il mercato quota moderatamente sopra i valori OMI. Situazione normale per offerte iniziali.')
        elif gap > -5:
            doc.add_paragraph('Il mercato è allineato ai valori OMI. Prezzi coerenti con le transazioni effettive.')
        else:
            doc.add_paragraph('Il mercato quota sotto i valori OMI. Possibili opportunità o necessità di rilancio del settore.')
        
        doc.add_paragraph()
        
        # Spazio per analisi AI futura

    
    # ========================================
    # SEZIONE ANALISI AI (se disponibile)
    # ========================================
    if analisi_ai and analisi_ai.get('success'):
        doc.add_page_break()
        doc.add_heading('🤖 Analisi AI con Claude', 1)
        
        p = doc.add_paragraph()
        run = p.add_run('Analisi generata automaticamente da Claude (Anthropic)')
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()
        
        # Gap Analysis
        if analisi_ai.get('gap_analysis'):
            gap = analisi_ai['gap_analysis']
            doc.add_heading('Gap OMI vs Mercato', 2)
            
            gap_table = doc.add_table(rows=4, cols=2)
            gap_table.style = 'Light Grid Accent 1'
            
            gap_data = [
                ['OMI Mediano', f"€{gap['omi_mediano']:,.0f}/m²"],
                ['Mercato Mediano', f"€{gap['mercato_mediano']:,.0f}/m²"],
                ['Differenza', f"€{gap['gap_assoluto']:,.0f}/m²"],
                ['Percentuale', f"{gap['gap_percentuale']:+.1f}%"]
            ]
            
            for i, (label, value) in enumerate(gap_data):
                gap_table.rows[i].cells[0].text = label
                gap_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                gap_table.rows[i].cells[1].text = value
            
            doc.add_paragraph()
        
        # Analisi testuale (semplice - niente markdown parsing)
        if analisi_ai.get('analisi_completa'):
            testo_analisi = analisi_ai['analisi_completa']
            
            # Dividi in paragrafi (ogni riga vuota = nuovo paragrafo)
            paragrafi = testo_analisi.split('\n\n')
            
            for para in paragrafi:
                para = para.strip()
                if not para:
                    continue
                
                # Se inizia con ##, è un titolo
                if para.startswith('## '):
                    titolo = para.replace('## ', '').strip()
                    doc.add_heading(titolo, 2)
                elif para.startswith('### '):
                    titolo = para.replace('### ', '').strip()
                    doc.add_heading(titolo, 3)
                else:
                    # Paragrafo normale
                    doc.add_paragraph(para)
        
        doc.add_paragraph()
        
        # Disclaimer
        p = doc.add_paragraph()
        run = p.add_run('Nota: ')
        run.font.bold = True
        run.font.size = Pt(9)
        run2 = p.add_run('Questa analisi è stata generata da intelligenza artificiale e deve essere considerata come supporto decisionale.')
        run2.font.size = Pt(9)
        run2.font.italic = True
        
        doc.add_paragraph()
    
        # ========================================
    # ========================================
    # SEZIONE MAPPA INTERATTIVA
    # ========================================
    if appartamenti and len(appartamenti) > 0:
        doc.add_page_break()
        doc.add_heading('🗺️ Mappa Appartamenti', 1)
        
        try:
            # Crea mappa
            mappa = crea_mappa_interattiva(
                lat_centro=lat,
                lon_centro=lon,
                via=via,
                comune=comune,
                raggio_km=raggio_km,
                appartamenti=appartamenti,
                stats_immobiliare=stats_immobiliare
            )
            
            # Salva mappa come HTML
            import os
            mappa_filename = f"mappa_{comune}_{now.strftime('%Y%m%d_%H%M%S')}.html"
            mappa_path = os.path.join(output_dir, mappa_filename)
            mappa.save(mappa_path)
            
            doc.add_paragraph(f'La mappa interattiva è stata salvata in: {mappa_filename}')
            doc.add_paragraph('Apri il file HTML per visualizzare la mappa con tutti gli appartamenti.')
            doc.add_paragraph()
            
            # Legenda
            doc.add_paragraph('Legenda:', style='Heading 3')
            doc.add_paragraph('📍 Pin rosso (casa): Centro ricerca', style='List Bullet')
            doc.add_paragraph('🏗️ Pin verde: Appartamenti economici (<-15% mediano)', style='List Bullet')
            doc.add_paragraph('🏗️ Pin blu: Appartamenti prezzo medio (±15% mediano)', style='List Bullet')
            doc.add_paragraph('🏗️ Pin arancione: Appartamenti alto prezzo (+15-35% mediano)', style='List Bullet')
            doc.add_paragraph('🏗️ Pin rosso: Appartamenti molto alto prezzo (>+35% mediano)', style='List Bullet')
            doc.add_paragraph()
            
            # Info coordinate
            from map_generator import get_mappa_statistiche
            map_stats = get_mappa_statistiche(appartamenti)
            
            doc.add_paragraph(f"Appartamenti visualizzati sulla mappa: {map_stats['con_coordinate']} su {map_stats['totale']}")
            
            if map_stats['senza_coordinate'] > 0:
                doc.add_paragraph(f"⚠️ {map_stats['senza_coordinate']} appartamenti non hanno coordinate e non sono visualizzati sulla mappa.")
            
            doc.add_paragraph()
            
        except Exception as e:
            doc.add_paragraph(f'⚠️ Impossibile generare la mappa: {str(e)}')
            print(f"[MAP][ERROR] Errore generazione mappa per report: {e}")
        
        doc.add_paragraph()
    
    # ========================================
    # FOOTER
    # ========================================
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.add_run(f'Report generato il {now.strftime("%d/%m/%Y alle %H:%M")}').italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('\nPlanet AI - Analisi Immobiliare').italic = True
    
    # Salva
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    filename = f"report_combinato_{comune}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)
    
    doc.save(filepath)
    
    print(f"   ✅ Report salvato: {filename}\n")
    
    return filepath