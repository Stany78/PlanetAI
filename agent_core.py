"""
AGENT 1.0 - Planet AI
---------------------

Flusso:

1. Input da tastiera:
      - Comune (es: Como)
      - Indirizzo (es: Via Borgovico 150)

2. Geocoding → (lat, lon)

3. OMI:
      - da coordinate → zona OMI (via KML, omi_utils)
      - da zona → valori €/mq (via CSV ufficiale)

4. Annunci:
      - scaricati in tempo reale da Immobiliare.it per il comune

   Se il portale non indica esplicitamente "nuovo", deduciamo gli annunci
   "nuovo (da prezzo)" prendendo quelli con prezzo/mq molto più alto
   della mediana dell'area.

5. Stima €/mq nuova costruzione:
      - stima da portali
      - stima da OMI
      - combinazione: 60% OMI + 40% portali (se entrambi disponibili)

6. Generazione report Word in BASE_DIR/reports con grafico OMI vs stima
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import List, Dict, Optional
import statistics
import os
import re
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from geopy.geocoders import Nominatim
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

from config import (
    BASE_DIR,
    MIN_COMPARABLE_NUOVO,
    SPREAD_NUOVO_MIN,
    SPREAD_NUOVO_MAX,
    FALLBACK_COORDINATE,
    DEBUG_MODE,
)

# Opzionali dal config
try:
    from config import FATTORE_DEFAULT_SU_OMI
except ImportError:
    FATTORE_DEFAULT_SU_OMI = 1.25  # nuovo ≈ +25% rispetto a OMI mediano

try:
    from config import OMI_DATA_INFO
except ImportError:
    OMI_DATA_INFO = "Dati OMI (semestre non specificato)"

from omi_utils import (
    get_quotazione_omi_da_coordinate,
    OMIQuotazione,
    warmup_omi_cache,
)


# ==============================
# Geocoding indirizzo → (lat, lon)
# ==============================

_geolocator = Nominatim(user_agent="planet_ai_omi_agent")


def geocode_indirizzo(comune: str, indirizzo: str) -> tuple[float, float]:
    """
    Geocoda 'indirizzo, comune, Italia' usando Nominatim.
    Se fallisce, usa FALLBACK_COORDINATE.
    """
    full_address = f"{indirizzo}, {comune}, Italia"
    if DEBUG_MODE:
        print(f"[GEO] Geocoding: {full_address}")

    try:
        loc = _geolocator.geocode(full_address)
        if loc is None:
            print("[GEO][WARN] Geocoding fallito, uso FALLBACK_COORDINATE.")
            return FALLBACK_COORDINATE
        return (loc.latitude, loc.longitude)
    except Exception as e:
        print(f"[GEO][ERROR] Geocoding errore: {e}. Uso FALLBACK_COORDINATE.")
        return FALLBACK_COORDINATE


# ==============================
# DATACLASS ANNUNCIO
# ==============================

@dataclass
class Annuncio:
    portale: str
    prezzo: float
    mq: float
    stato: Optional[str] = None
    anno_costruzione: Optional[int] = None
    classe_energetica: Optional[str] = None


# ==============================
# SCRAPER IMMOBILIARE.IT
# ==============================

def _slugify_nome_comune(nome: str) -> str:
    """
    Converte "Como" -> "como",
    "San Giovanni la Punta" -> "san-giovanni-la-punta"
    per costruire l'URL https://www.immobiliare.it/vendita-case/<slug>/
    """
    s = nome.lower().strip()

    # gestione basilare accenti
    for a, b in [
        ("à", "a"), ("è", "e"), ("é", "e"),
        ("ì", "i"), ("ò", "o"), ("ù", "u")
    ]:
        s = s.replace(a, b)

    s = re.sub(r"[^a-z0-9]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    return s or nome.lower()


def _scrape_immobiliare_it(comune: str, max_annunci: int = 60) -> List[Annuncio]:
    """
    Scarica annunci da Immobiliare.it per un intero comune.

    Approccio robusto a testo:
    - va su: https://www.immobiliare.it/vendita-case/<slug-comune>/
    - usa BeautifulSoup per estrarre il testo
    - trova pattern:
         € 245.000 ... 43 m²
    - nella finestra di testo attorno a prezzo+mq cerca parole chiave
      per classificare nuovo / usato / da ristrutturare.
    """
    slug = _slugify_nome_comune(comune)
    url = f"https://www.immobiliare.it/vendita-case/{slug}/"

    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; PlanetAI/1.0; +https://example.com)"
    }

    if DEBUG_MODE:
        print(f"[SCRAPER] Richiedo pagina Immobiliare.it: {url}")

    try:
        resp = requests.get(url, headers=headers, timeout=20)
    except Exception as e:
        print(f"[SCRAPER][ERROR] Errore richiesta Immobiliare.it: {e}")
        return []

    if resp.status_code != 200:
        print(f"[SCRAPER][WARN] Status code {resp.status_code} da Immobiliare.it")
        return []

    soup = BeautifulSoup(resp.text, "html.parser")
    text = soup.get_text(separator="\n")

    # Pattern tipo:
    #   € 245.000  ...   43 m²
    pattern = re.compile(r"€\s*([\d\.\s]+).*?(\d+)\s*m²", re.DOTALL)

    parole_nuovo = [
        "nuova costruzione",
        "nuovo",
        "nuovissimo",
        "di nuova costruzione",
        "in costruzione",
        "recentissima costruzione",
        "mai abitato",
    ]
    parole_ristr = [
        "ristrutturare",
        "da ristrutturare",
        "completamente da ristrutturare",
    ]

    annunci: List[Annuncio] = []

    for m in pattern.finditer(text):
        prezzo_raw = m.group(1)
        mq_raw = m.group(2)

        # pulizia prezzo "245.000" -> 245000
        prezzo_str = re.sub(r"[^\d]", "", prezzo_raw)
        if not prezzo_str:
            continue

        try:
            prezzo = float(prezzo_str)
            mq = float(mq_raw.replace(",", "."))
        except ValueError:
            continue

        # finestra di contesto PIÙ AMPIA attorno a prezzo/mq
        start = max(0, m.start() - 800)
        end = min(len(text), m.end() + 800)
        window = text[start:end].lower()

        stato = "usato"
        if any(k in window for k in parole_nuovo):
            stato = "nuova costruzione"
        elif any(k in window for k in parole_ristr):
            stato = "da ristrutturare"

        annunci.append(
            Annuncio(
                portale="immobiliare.it",
                prezzo=prezzo,
                mq=mq,
                stato=stato,
                anno_costruzione=None,
            )
        )

        if len(annunci) >= max_annunci:
            break

    if DEBUG_MODE:
        print(f"[SCRAPER] immobiliare.it: trovati {len(annunci)} annunci per {comune}")

    return annunci


# ==============================
# "Nuovo" dedotto dai prezzi
# ==============================

def rietichetta_nuovi_da_prezzi(
    annunci: List[Annuncio],
    soglia_spread: float = 0.25,
    minimo_nuovo: int = 3,
) -> None:
    """
    Se non ci sono annunci marcati come 'nuovi', ma ci sono abbastanza dati,
    deduce i 'nuovi' dai prezzi/mq più alti.

    Logica:
      1. Calcola la mediana dei prezzi/mq (escludendo 'da ristrutturare').
      2. Se prezzo/mq di un annuncio 'usato' è > mediana * (1 + soglia_spread),
         lo considera candidato 'nuovo'.
      3. Etichetta come 'nuovo (da prezzo)' i primi 'minimo_nuovo' candidati.

    Questo serve per non avere mai completamente vuoto il segmento "nuovo"
    quando il portale non scrive "nuova costruzione" nel testo.
    """
    prezzi_mq = []
    for a in annunci:
        if not a.prezzo or not a.mq:
            continue
        if a.stato and "ristrutturare" in a.stato.lower():
            continue
        prezzi_mq.append(a.prezzo / a.mq)

    if len(prezzi_mq) < 5:
        # troppo pochi annunci per fare statistica
        return

    mediana = statistics.median(prezzi_mq)
    soglia = mediana * (1.0 + soglia_spread)

    candidati = []
    for a in annunci:
        if not a.prezzo or not a.mq:
            continue
        if a.stato and "ristrutturare" in a.stato.lower():
            continue
        # se già nuovo, non toccare
        if a.stato and "nuovo" in a.stato.lower():
            continue
        prezzo_mq = a.prezzo / a.mq
        if prezzo_mq >= soglia:
            candidati.append((a, prezzo_mq))

    if not candidati:
        return

    # ordina per prezzo/mq decrescente e prende i top N
    candidati.sort(key=lambda x: x[1], reverse=True)

    for a, _ in candidati[:minimo_nuovo]:
        a.stato = "nuovo (da prezzo)"


# ==============================
# ANNUNCI ZONA (aggregatore portali)
# ==============================

def scarica_annunci_zona(
    lat: float,
    lon: float,
    comune: Optional[str] = None,
    indirizzo: Optional[str] = None,
    raggio_km: float = 2.0
) -> List[Annuncio]:
    """
    Aggregatore di annunci per la zona.
    Per ora:
      - usa SOLO Immobiliare.it, a livello di COMUNE (non di singola via).

    In futuro potremo usare lat/lon e raggio per ricerche su mappa
    e aggiungere altri portali (idealista, casa.it, ecc).
    """
    annunci: List[Annuncio] = []

    # 1) Immobiliare.it
    if comune:
        try:
            annunci_imm = _scrape_immobiliare_it(comune)
            annunci.extend(annunci_imm)
        except Exception as e:
            print(f"[SCRAPER][ERROR] Immobiliare.it ha dato errore: {e}")

    # Se non c'è nessun annuncio marcato come "nuovo", proviamo a dedurli dai prezzi
    if annunci and not any(a.stato and "nuovo" in a.stato.lower() for a in annunci):
        if DEBUG_MODE:
            print("[SCRAPER] Nessun 'nuovo' esplicito: provo a dedurre dai prezzi/mq più alti.")
        rietichetta_nuovi_da_prezzi(annunci)

    if DEBUG_MODE:
        print(f"[SCRAPER] Totale annunci raccolti: {len(annunci)}")

    return annunci


# ==============================
# LOGICA DI CLASSIFICAZIONE ANNUNCI
# ==============================

def classifica_annuncio(a: Annuncio) -> str:
    stato = (a.stato or "").lower()

    if "nuova costruzione" in stato or "nuovo" in stato or "nuovo (da prezzo)" in stato:
        return "nuovo"

    if a.anno_costruzione is not None and a.anno_costruzione >= 2018:
        return "nuovo"

    if "ristrutturare" in stato:
        return "ristrutturare"

    return "usato"


def calcola_statistiche(prezzi_mq: List[float]) -> Optional[Dict[str, float]]:
    if not prezzi_mq:
        return None

    prezzi_mq = sorted(prezzi_mq)
    n = len(prezzi_mq)

    mediana = statistics.median(prezzi_mq)
    p25 = prezzi_mq[int(0.25 * (n - 1))]
    p75 = prezzi_mq[int(0.75 * (n - 1))]

    return {"n": n, "mediana": mediana, "p25": p25, "p75": p75}


# ==============================
# FUNZIONE PRINCIPALE DI STIMA
# ==============================

def stima_prezzo_zona_nuova_costruzione(comune: str, indirizzo: str) -> Dict:
    """
    1. Geocoding
    2. OMI da coordinate
    3. Annunci da portale (Immobiliare.it)
    4. Stima:
         - da portali
         - da OMI
         - combinazione 60% OMI + 40% portali
    """

    comune_norm = comune.strip()
    indirizzo_norm = indirizzo.strip()

    # 1️⃣ Geocoding
    lat, lon = geocode_indirizzo(comune_norm, indirizzo_norm)

    # 2️⃣ OMI da coordinate
    zona_omi: Optional[OMIQuotazione] = get_quotazione_omi_da_coordinate(lat, lon)
    if zona_omi:
        print(
            f"[OMI] Zona OMI: {zona_omi.comune} {zona_omi.zona_codice} "
            f"(min={zona_omi.val_min_mq}, med={zona_omi.val_med_mq}, max={zona_omi.val_max_mq} €/mq)"
        )
    else:
        print("[OMI] Nessuna zona OMI trovata per queste coordinate (KML/CSV non matchano).")

    # 3️⃣ Annunci reali (Immobiliare.it sul comune)
    annunci = scarica_annunci_zona(lat, lon, comune=comune_norm, indirizzo=indirizzo_norm)
    prezzi_nuovo: List[float] = []
    prezzi_usato: List[float] = []

    for a in annunci:
        if not a.prezzo or not a.mq:
            continue

        gruppo = classifica_annuncio(a)
        prezzo_mq = a.prezzo / a.mq

        if gruppo == "nuovo":
            prezzi_nuovo.append(prezzo_mq)
        elif gruppo == "usato":
            prezzi_usato.append(prezzo_mq)
        # 'ristrutturare' ignorato per la stima del nuovo

    stat_nuovo = calcola_statistiche(prezzi_nuovo)
    stat_usato = calcola_statistiche(prezzi_usato)

    if DEBUG_MODE:
        print(f"[STATS] NUOVO (da annunci portali): {stat_nuovo}")
        print(f"[STATS] USATO (da annunci portali): {stat_usato}")

    # 4️⃣ Stima da portali (solo annunci)
    stima_portali = {"prudente": None, "centrale": None, "aggressivo": None}
    spread = None

    if stat_nuovo and stat_nuovo["n"] >= MIN_COMPARABLE_NUOVO:
        stima_portali["centrale"] = stat_nuovo["mediana"]
        stima_portali["prudente"] = stat_nuovo["p25"]
        stima_portali["aggressivo"] = stat_nuovo["p75"]
        if stat_usato:
            spread = (stat_nuovo["mediana"] / stat_usato["mediana"]) - 1.0

    elif stat_nuovo and stat_nuovo["n"] > 0:
        stima_portali["centrale"] = stat_nuovo["mediana"]
        stima_portali["prudente"] = stat_nuovo["p25"]
        stima_portali["aggressivo"] = stat_nuovo["p75"]
        if stat_usato:
            spread = (stat_nuovo["mediana"] / stat_usato["mediana"]) - 1.0

    elif stat_usato:
        spread = (SPREAD_NUOVO_MIN + SPREAD_NUOVO_MAX) / 2
        stima_portali["centrale"] = stat_usato["mediana"] * (1 + spread)
        stima_portali["prudente"] = stat_usato["p25"] * (1 + SPREAD_NUOVO_MIN)
        stima_portali["aggressivo"] = stat_usato["p75"] * (1 + SPREAD_NUOVO_MAX)

    if DEBUG_MODE:
        print(f"[STIMA PORTALI] {stima_portali}")

    # 5️⃣ Stima da OMI (se disponibile)
    stima_omi = {"prudente": None, "centrale": None, "aggressivo": None}
    if zona_omi and zona_omi.val_med_mq is not None:
        stima_omi["centrale"] = zona_omi.val_med_mq * FATTORE_DEFAULT_SU_OMI
        stima_omi["prudente"] = (zona_omi.val_min_mq or zona_omi.val_med_mq) * 1.2
        stima_omi["aggressivo"] = (zona_omi.val_max_mq or zona_omi.val_med_mq) * 1.4

    if DEBUG_MODE and stima_omi["centrale"] is not None:
        print(f"[STIMA OMI] {stima_omi}")

    # 6️⃣ Combinazione OMI + portali
    prudente = centrale = aggressivo = None
    note_base = ""

    portali_ok = stima_portali["centrale"] is not None
    omi_ok = stima_omi["centrale"] is not None

    if portali_ok and omi_ok:
        peso_omi = 0.6
        peso_portali = 0.4

        prudente = stima_omi["prudente"] * peso_omi + stima_portali["prudente"] * peso_portali
        centrale = stima_omi["centrale"] * peso_omi + stima_portali["centrale"] * peso_portali
        aggressivo = stima_omi["aggressivo"] * peso_omi + stima_portali["aggressivo"] * peso_portali

        note_base = "Stima combinata: 60% valori OMI + 40% annunci di mercato (Immobiliare.it)."

    elif omi_ok:
        prudente = stima_omi["prudente"]
        centrale = stima_omi["centrale"]
        aggressivo = stima_omi["aggressivo"]
        note_base = "Stima basata unicamente su valori OMI (nessun dato portali utilizzabile)."

    elif portali_ok:
        prudente = stima_portali["prudente"]
        centrale = stima_portali["centrale"]
        aggressivo = stima_portali["aggressivo"]
        note_base = "Stima basata unicamente su annunci di mercato (Immobiliare.it), OMI non disponibile."

    else:
        prudente = 2500.0
        centrale = 3000.0
        aggressivo = 3500.0
        note_base = "Nessun dato OMI né portali: stima di fallback generica."

    note: List[str] = [note_base]

    if stat_nuovo:
        note.append(f"Annunci 'nuovo' considerati: n={stat_nuovo['n']} (Immobiliare.it, inclusi eventuali 'nuovo (da prezzo)').")
    else:
        note.append("Nessun annuncio 'nuovo' utilizzabile nemmeno dopo l'analisi dei prezzi.")

    if stat_usato:
        note.append(f"Annunci 'usato' considerati: n={stat_usato['n']} (Immobiliare.it).")
    else:
        note.append("Nessun annuncio 'usato' utilizzabile (Immobiliare.it).")

    if zona_omi is None:
        note.append("Zona OMI non determinata (verificare copertura KML/CSV).")

    zona_omi_dict = None
    if zona_omi:
        zona_omi_dict = {
            "codice": zona_omi.zona_codice,
            "descrizione": zona_omi.zona_descrizione,
            "comune": zona_omi.comune,
            "val_min": zona_omi.val_min_mq,
            "val_med": zona_omi.val_med_mq,
            "val_max": zona_omi.val_max_mq,
        }

    result = {
        "input": {
            "comune": comune_norm,
            "indirizzo": indirizzo_norm,
            "coordinate": {"lat": lat, "lon": lon},
        },
        "zona_omi": zona_omi_dict,
        "usato": stat_usato or {},
        "nuovo": stat_nuovo or {},
        "spread_nuovo_vs_usato": spread,
        "stima_nuova_costruzione": {
            "prudente": prudente,
            "centrale": centrale,
            "aggressivo": aggressivo,
        },
        "note": note,
    }

    if DEBUG_MODE:
        print("\n[AGENT RESULT - STIMA FINALE]")
        for k, v in result["stima_nuova_costruzione"].items():
            print(f"  {k.upper():<10}: {v:,.0f} €/mq".replace(",", "."))

    return result


# ==============================
# GRAFICO OMI vs STIMA (solo salvato)
# ==============================

def crea_grafico_omi_vs_stima(
    zona_omi: Optional[Dict],
    stima: Dict[str, float],
    output_dir: str,
) -> Optional[str]:
    """
    Crea un grafico a barre con:
      - valori OMI (min/med/max) se disponibili
      - stima nuova costruzione (prudente/centrale/aggressivo)
    Lo salva come PNG e NON apre la finestra grafica.
    """
    try:
        os.makedirs(output_dir, exist_ok=True)

        labels = []
        values = []

        # Valori OMI se disponibili
        if zona_omi and zona_omi.get("val_med") is not None:
            labels.extend(["OMI min", "OMI med", "OMI max"])
            values.extend([
                zona_omi["val_min"],
                zona_omi["val_med"],
                zona_omi["val_max"],
            ])

        # Stima nuova costruzione
        labels.extend(["Stima prud.", "Stima centr.", "Stima aggr."])
        values.extend([
            stima["prudente"],
            stima["centrale"],
            stima["aggressivo"],
        ])

        plt.figure(figsize=(8, 4))
        x = range(len(labels))
        bars = plt.bar(x, values)

        plt.xticks(x, labels, rotation=20, ha="right")
        plt.ylabel("€/mq")
        plt.title("Confronto valori OMI vs stima nuova costruzione")
        plt.grid(axis="y", linestyle="--", alpha=0.5)

        # etichette sopra le barre
        for bar, val in zip(bars, values):
            height = bar.get_height()
            plt.text(
                bar.get_x() + bar.get_width() / 2,
                height,
                f"{val:,.0f}".replace(",", "."),
                ha="center",
                va="bottom",
                fontsize=8,
            )

        plt.tight_layout()

        fname = f"grafico_omi_stima_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        filepath = os.path.join(output_dir, fname)

        plt.savefig(filepath, dpi=150)
        plt.close()

        if DEBUG_MODE:
            print(f"[GRAFICO] Creato grafico OMI vs stima: {filepath}")

        return filepath

    except Exception as e:
        print(f"[GRAFICO][ERROR] Impossibile creare il grafico: {e}")
        return None


# ==============================
# GENERAZIONE REPORT WORD
# ==============================

def genera_report_word(risultato: Dict, output_dir: Optional[str] = None) -> Optional[str]:
    """
    Crea un report in formato Word (.docx) con:
      - dati di input
      - zona OMI + valori OMI (se disponibili) + data OMI
      - statistiche annunci
      - stima €/mq nuova costruzione
      - note
      - grafico OMI vs stima
    Ritorna il percorso del file o None se fallisce.
    """
    try:
        if output_dir is None:
            output_dir = os.path.join(BASE_DIR, "reports")

        os.makedirs(output_dir, exist_ok=True)

        comune = risultato["input"]["comune"]
        indirizzo = risultato["input"]["indirizzo"]
        coord = risultato["input"]["coordinate"]
        stima = risultato["stima_nuova_costruzione"]
        usato = risultato.get("usato")
        nuovo = risultato.get("nuovo")
        zona_omi = risultato.get("zona_omi")
        note = risultato.get("note", [])

        zona_codice = zona_omi["codice"] if zona_omi else "N/D"

        safe_comune = comune.replace(" ", "_")
        safe_ind = indirizzo.replace(" ", "_").replace("/", "-")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Report_{safe_comune}_{safe_ind}_{timestamp}.docx"
        filepath = os.path.join(output_dir, filename)

        doc = Document()

        # Titolo
        titolo = doc.add_heading("Report di Stima Immobiliare – Planet AI", level=1)
        titolo.alignment = 1  # centrato

        # Dati base
        doc.add_paragraph(f"Data report: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Comune: {comune}")
        doc.add_paragraph(f"Indirizzo: {indirizzo}")
        doc.add_paragraph(f"Zona OMI: {zona_codice}")
        doc.add_paragraph(f"Coordinate: lat={coord['lat']:.6f}, lon={coord['lon']:.6f}")

        doc.add_paragraph("")

        # Valori OMI
        doc.add_heading("Valori OMI (fonte Agenzia delle Entrate)", level=2)
        doc.add_paragraph(f"Dati OMI: {OMI_DATA_INFO}")

        if zona_omi and zona_omi["val_med"] is not None:
            tab_omi = doc.add_table(rows=2, cols=4)
            hdr = tab_omi.rows[0].cells
            hdr[0].text = "Comune"
            hdr[1].text = "Zona"
            hdr[2].text = "Min €/mq"
            hdr[3].text = "Med / Max €/mq"

            row = tab_omi.rows[1].cells
            row[0].text = zona_omi["comune"]
            row[1].text = zona_omi["codice"]
            row[2].text = f"{zona_omi['val_min']:.0f}"
            row[3].text = f"{zona_omi['val_med']:.0f} / {zona_omi['val_max']:.0f}"
        else:
            doc.add_paragraph("Valori OMI non disponibili per questa zona (nessun match nel CSV).")

        doc.add_paragraph("")

        # Statistiche annunci
        doc.add_heading("Statistiche annunci (Immobiliare.it)", level=2)

        tab = doc.add_table(rows=1, cols=5)
        hdr_cells = tab.rows[0].cells
        hdr_cells[0].text = "Tipo"
        hdr_cells[1].text = "Mediana €/mq"
        hdr_cells[2].text = "P25"
        hdr_cells[3].text = "P75"
        hdr_cells[4].text = "N annunci"

        def add_row(label: str, stats: Optional[Dict[str, float]]):
            row_cells = tab.add_row().cells
            row_cells[0].text = label
            if stats:
                row_cells[1].text = f"{stats['mediana']:.0f}"
                row_cells[2].text = f"{stats['p25']:.0f}"
                row_cells[3].text = f"{stats['p75']:.0f}"
                row_cells[4].text = f"{stats['n']}"
            else:
                for i in range(1, 5):
                    row_cells[i].text = "-"

        add_row("Usato", usato)
        add_row("Nuovo", nuovo)

        doc.add_paragraph("")

        # Stima nuova costruzione
        doc.add_heading("Stima nuova costruzione (€ / mq)", level=2)
        for k, v in stima.items():
            if v is not None:
                doc.add_paragraph(f"{k.capitalize()}: {v:,.0f} €/mq".replace(",", "."))

        doc.add_paragraph("")

        # Note
        if note:
            doc.add_heading("Note", level=2)
            for n in note:
                doc.add_paragraph(f"– {n}")

        # Grafico OMI vs Stima
        doc.add_paragraph("")
        doc.add_heading("Confronto grafico OMI vs stima", level=2)

        grafico_path = crea_grafico_omi_vs_stima(
            zona_omi=zona_omi,
            stima=stima,
            output_dir=output_dir,
        )

        if grafico_path:
            doc.add_picture(grafico_path, width=Inches(5.5))
        else:
            doc.add_paragraph("Grafico non disponibile (errore in generazione).")

        doc.save(filepath)

        print(f"[REPORT] Report Word generato: {filepath}")

        try:
            if os.name == "nt":  # Windows
                os.startfile(filepath)
            else:
                print("Apri il file manualmente da questo percorso.")
        except Exception as e:
            print(f"[REPORT][WARN] Non riesco ad aprire Word automaticamente: {e}")

        return filepath

    except Exception as e:
        print(f"[REPORT][ERROR] Impossibile generare il report Word: {e}")
        return None


# ==============================
# MAIN - INPUT MANUALE
# ==============================

if __name__ == "__main__":
    # Pre-carica KML + CSV (cache)
    warmup_omi_cache()

    print("=== AGENT 1.0 - Stima nuova costruzione ===")
    comune = input("Inserisci il COMUNE (es: Como): ").strip()
    indirizzo = input("Inserisci l'INDIRIZZO (es: Via Borgovico 150): ").strip()

    risultato = stima_prezzo_zona_nuova_costruzione(comune, indirizzo)

    print("\n[OUTPUT CONSOLE]")
    zona = risultato.get("zona_omi")
    if zona:
        print(f"Zona OMI: {zona['comune']} {zona['codice']}")
        if zona["val_med"] is not None:
            print(
                f"Valori OMI €/mq - min: {zona['val_min']:.0f}, "
                f"med: {zona['val_med']:.0f}, max: {zona['val_max']:.0f}"
            )
        else:
            print("Valori OMI €/mq: non disponibili nel CSV per questa zona.")
    else:
        print("Zona OMI: non determinata")

    print("Stima nuova costruzione (€ / mq):")
    for k, v in risultato["stima_nuova_costruzione"].items():
        print(f"  {k.capitalize():<10}: {v:,.0f} €/mq".replace(",", "."))

    genera_report_word(risultato)